# -*- coding: utf-8 -*-
"""
Document Generator - Generate PDF, Word, and Excel documents
Supports templates, formatting, and dynamic content insertion
"""

import os
import logging
import re
from typing import Dict, List, Optional, Any, Union
from pathlib import Path
from datetime import datetime
import json

logger = logging.getLogger(__name__)

# PDF generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger.warning("reportlab not available. PDF generation will be limited.")

# Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Word document generation will be limited.")

# Excel generation
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl not available. Excel generation will be limited.")

# Markdown to HTML/PDF
try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    logger.warning("markdown not available. Markdown conversion will be limited.")

# Barcode and QR code generation
try:
    import qrcode
    from qrcode.image.pil import PilImage
    QRCODE_AVAILABLE = True
except ImportError:
    QRCODE_AVAILABLE = False
    logger.warning("qrcode not available. QR code generation will be limited.")

try:
    import barcode
    from barcode.writer import ImageWriter
    BARCODE_AVAILABLE = True
except ImportError:
    BARCODE_AVAILABLE = False
    logger.warning("python-barcode not available. Barcode generation will be limited.")

# PSD template processing
try:
    from psd_tools import PSDImage
    PSD_TOOLS_AVAILABLE = True
except ImportError:
    PSD_TOOLS_AVAILABLE = False
    logger.warning("psd-tools not available. PSD template processing will be limited.")


class DocumentGenerator:
    """Generate PDF, Word, and Excel documents with formatting and templates"""
    
    def __init__(self, output_dir: str = "generated_documents"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Check availability
        self.pdf_available = REPORTLAB_AVAILABLE
        self.docx_available = DOCX_AVAILABLE
        self.xlsx_available = OPENPYXL_AVAILABLE
        self.markdown_available = MARKDOWN_AVAILABLE
        self.qrcode_available = QRCODE_AVAILABLE
        self.barcode_available = BARCODE_AVAILABLE
        
        logger.info(f"Document Generator initialized - PDF: {self.pdf_available}, Word: {self.docx_available}, Excel: {self.xlsx_available}, QR: {self.qrcode_available}, Barcode: {self.barcode_available}")
    
    def generate_pdf(self, 
                     content: Union[str, Dict],
                     filename: str = None,
                     title: str = None,
                     author: str = "SMG-Forcer Bot",
                     page_size: str = "letter",
                     **kwargs) -> Optional[str]:
        """
        Generate PDF document
        
        Args:
            content: Text content, markdown, or structured dict
            filename: Output filename (auto-generated if None)
            title: Document title
            author: Document author
            page_size: 'letter' or 'A4'
            **kwargs: Additional formatting options
        
        Returns:
            Path to generated PDF file or None if failed
        """
        if not self.pdf_available:
            logger.error("reportlab not available for PDF generation")
            return None
        
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"document_{timestamp}.pdf"
            
            if not filename.endswith('.pdf'):
                filename += '.pdf'
            
            filepath = self.output_dir / filename
            
            # Select page size
            size = A4 if page_size.upper() == 'A4' else letter
            
            # Create PDF document
            doc = SimpleDocTemplate(str(filepath), pagesize=size)
            story = []
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Add title
            if title:
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=24,
                    textColor=colors.HexColor('#1a1a1a'),
                    spaceAfter=30,
                    alignment=TA_CENTER
                )
                story.append(Paragraph(title, title_style))
                story.append(Spacer(1, 0.2*inch))
            
            # Process content
            if isinstance(content, dict):
                # Structured content
                story.extend(self._process_structured_content(content, styles))
            elif isinstance(content, str):
                # Text or markdown content
                if self.markdown_available and self._is_markdown(content):
                    # Convert markdown to HTML then to PDF
                    html_content = markdown.markdown(content, extensions=['extra', 'codehilite'])
                    story.extend(self._html_to_paragraphs(html_content, styles))
                else:
                    # Plain text
                    paragraphs = content.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            story.append(Paragraph(para.strip(), styles['Normal']))
                            story.append(Spacer(1, 0.1*inch))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"PDF generated: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generating PDF: {e}", exc_info=True)
            return None
    
    def generate_word(self,
                     content: Union[str, Dict],
                     filename: str = None,
                     title: str = None,
                     **kwargs) -> Optional[str]:
        """
        Generate Word document (.docx)
        
        Args:
            content: Text content, markdown, or structured dict
            filename: Output filename (auto-generated if None)
            title: Document title
            **kwargs: Additional formatting options
        
        Returns:
            Path to generated Word file or None if failed
        """
        if not self.docx_available:
            logger.error("python-docx not available for Word generation")
            return None
        
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"document_{timestamp}.docx"
            
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            filepath = self.output_dir / filename
            
            # Create document
            doc = Document()
            
            # Add title
            if title:
                title_para = doc.add_heading(title, 0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Process content
            if isinstance(content, dict):
                self._add_structured_content_to_word(doc, content)
            elif isinstance(content, str):
                if self.markdown_available and self._is_markdown(content):
                    # Convert markdown to Word
                    self._markdown_to_word(doc, content)
                else:
                    # Plain text
                    paragraphs = content.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
            
            # Save document
            doc.save(str(filepath))
            
            logger.info(f"Word document generated: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generating Word document: {e}", exc_info=True)
            return None
    
    def generate_excel(self,
                      data: Union[List[List[Any]], Dict[str, List[List[Any]]]],
                      filename: str = None,
                      sheet_name: str = "Sheet1",
                      headers: List[str] = None,
                      **kwargs) -> Optional[str]:
        """
        Generate Excel spreadsheet (.xlsx)
        
        Args:
            data: List of rows (list of lists) or dict with sheet names as keys
            filename: Output filename (auto-generated if None)
            sheet_name: Name for single sheet (if data is list)
            headers: Column headers (optional)
            **kwargs: Additional formatting options
        
        Returns:
            Path to generated Excel file or None if failed
        """
        if not self.xlsx_available:
            logger.error("openpyxl not available for Excel generation")
            return None
        
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"spreadsheet_{timestamp}.xlsx"
            
            if not filename.endswith('.xlsx'):
                filename += '.xlsx'
            
            filepath = self.output_dir / filename
            
            # Create workbook
            wb = Workbook()
            
            # Remove default sheet if we have multiple sheets
            if isinstance(data, dict) and len(data) > 1:
                wb.remove(wb.active)
            
            # Process data
            if isinstance(data, dict):
                # Multiple sheets
                for sheet_name, sheet_data in data.items():
                    ws = wb.create_sheet(title=sheet_name)
                    self._add_data_to_excel_sheet(ws, sheet_data, headers)
            else:
                # Single sheet
                ws = wb.active
                ws.title = sheet_name
                self._add_data_to_excel_sheet(ws, data, headers)
            
            # Save workbook
            wb.save(str(filepath))
            
            logger.info(f"Excel spreadsheet generated: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generating Excel spreadsheet: {e}", exc_info=True)
            return None
    
    def generate_from_template(self,
                              template_data: Dict,
                              doc_type: str = "pdf",
                              filename: str = None,
                              variables: Dict[str, Any] = None,
                              template_name: str = None) -> Optional[str]:
        """
        Generate document from template data
        
        Args:
            template_data: Template structure (from database)
            doc_type: 'pdf', 'docx', or 'xlsx'
            filename: Output filename
            variables: Variables to fill in template placeholders
            template_name: Name of template to use (for PSD templates)
        
        Returns:
            Path to generated document or None if failed
        """
        try:
            # Check if template_name is provided and we have a PSD template
            if template_name:
                try:
                    from template_processor import get_template_processor
                    processor = get_template_processor()
                    template_info = processor.get_template_info(template_name)
                    
                    if template_info and template_info.get('file_path'):
                        # Use PSD template
                        return self._generate_from_psd_template(template_info, variables or {}, filename)
                except Exception as e:
                    logger.warning(f"Could not use PSD template: {e}")
            
            # Fill template variables
            content = self._fill_template(template_data, variables or {})
            
            # Generate based on type
            if doc_type.lower() == 'pdf':
                return self.generate_pdf(content, filename=filename, **template_data.get('options', {}))
            elif doc_type.lower() == 'docx':
                return self.generate_word(content, filename=filename, **template_data.get('options', {}))
            elif doc_type.lower() == 'xlsx':
                return self.generate_excel(content, filename=filename, **template_data.get('options', {}))
            else:
                logger.error(f"Unknown document type: {doc_type}")
                return None
                
        except Exception as e:
            logger.error(f"Error generating from template: {e}", exc_info=True)
            return None
    
    def _process_structured_content(self, content: Dict, styles) -> List:
        """Process structured content dict for PDF"""
        story = []
        
        for section in content.get('sections', []):
            # Section title
            if 'title' in section:
                story.append(Paragraph(section['title'], styles['Heading2']))
                story.append(Spacer(1, 0.1*inch))
            
            # Section content
            if 'content' in section:
                if isinstance(section['content'], list):
                    for item in section['content']:
                        story.append(Paragraph(str(item), styles['Normal']))
                        story.append(Spacer(1, 0.05*inch))
                else:
                    story.append(Paragraph(str(section['content']), styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            # Tables
            if 'table' in section:
                table_data = section['table']
                if isinstance(table_data, list) and len(table_data) > 0:
                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 14),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 0.2*inch))
        
        return story
    
    def _add_structured_content_to_word(self, doc: Document, content: Dict):
        """Add structured content to Word document"""
        for section in content.get('sections', []):
            if 'title' in section:
                doc.add_heading(section['title'], level=2)
            
            if 'content' in section:
                if isinstance(section['content'], list):
                    for item in section['content']:
                        doc.add_paragraph(str(item))
                else:
                    doc.add_paragraph(str(section['content']))
            
            if 'table' in section:
                table_data = section['table']
                if isinstance(table_data, list) and len(table_data) > 0:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    for i, row_data in enumerate(table_data):
                        for j, cell_data in enumerate(row_data):
                            table.rows[i].cells[j].text = str(cell_data)
    
    def _markdown_to_word(self, doc: Document, markdown_text: str):
        """Convert markdown to Word document"""
        html = markdown.markdown(markdown_text, extensions=['extra', 'codehilite'])
        
        # Simple conversion - split by paragraphs
        paragraphs = markdown_text.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Check for headers
            if para.startswith('#'):
                level = len(para) - len(para.lstrip('#'))
                text = para.lstrip('#').strip()
                doc.add_heading(text, level=min(level, 9))
            elif para.startswith('-') or para.startswith('*'):
                # List item
                doc.add_paragraph(para.lstrip('-*').strip(), style='List Bullet')
            else:
                doc.add_paragraph(para)
    
    def _add_data_to_excel_sheet(self, ws, data: List[List[Any]], headers: List[str] = None):
        """Add data to Excel worksheet with formatting"""
        # Add headers if provided
        if headers:
            header_row = 1
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=header_row, column=col_idx, value=header)
                cell.font = Font(bold=True, size=12)
                cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                cell.font = Font(bold=True, color="FFFFFF", size=12)
                cell.alignment = Alignment(horizontal="center", vertical="center")
            
            # Start data from row 2
            start_row = 2
        else:
            start_row = 1
        
        # Add data
        for row_idx, row_data in enumerate(data, start=start_row):
            for col_idx, cell_value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=cell_value)
                cell.alignment = Alignment(horizontal="left", vertical="center")
        
        # Auto-adjust column widths
        for col_idx in range(1, max(len(row) for row in data) + 1):
            column_letter = get_column_letter(col_idx)
            max_length = 0
            for row in data:
                if col_idx <= len(row):
                    try:
                        if len(str(row[col_idx - 1])) > max_length:
                            max_length = len(str(row[col_idx - 1]))
                    except:
                        pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
    
    def _html_to_paragraphs(self, html: str, styles) -> List:
        """Convert HTML to ReportLab paragraphs (simplified)"""
        story = []
        # Simple HTML parsing - split by tags
        # Remove HTML tags for now (can be enhanced)
        text = re.sub(r'<[^>]+>', '', html)
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        return story
    
    def _is_markdown(self, text: str) -> bool:
        """Check if text appears to be markdown"""
        markdown_patterns = [
            r'^#{1,6}\s',  # Headers
            r'^\*\s',      # Bullet lists
            r'^\d+\.\s',   # Numbered lists
            r'\*\*.*\*\*', # Bold
            r'`.*`',       # Code
            r'\[.*\]\(.*\)', # Links
        ]
        return any(re.search(pattern, text, re.MULTILINE) for pattern in markdown_patterns)
    
    def _fill_template(self, template_data: Dict, variables: Dict[str, Any]) -> Union[str, Dict]:
        """Fill template placeholders with variables"""
        import json
        
        # Convert template_data to string for replacement
        template_str = json.dumps(template_data) if isinstance(template_data, dict) else str(template_data)
        
        # Replace placeholders {variable_name}
        for key, value in variables.items():
            placeholder = f"{{{key}}}"
            template_str = template_str.replace(placeholder, str(value))
        
        # Also replace common placeholders
        common_vars = {
            '{date}': datetime.now().strftime('%Y-%m-%d'),
            '{time}': datetime.now().strftime('%H:%M:%S'),
            '{datetime}': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        }
        for placeholder, value in common_vars.items():
            template_str = template_str.replace(placeholder, value)
        
        # Try to parse back to dict if it was a dict
        try:
            return json.loads(template_str)
        except:
            return template_str
    
    def generate_qr_code(self,
                        data: str,
                        filename: str = None,
                        size: int = 10,
                        border: int = 4,
                        error_correction: str = 'M',
                        box_size: int = 10) -> Optional[str]:
        """
        Generate QR code image
        
        Args:
            data: Data to encode in QR code
            filename: Output filename (auto-generated if None)
            size: QR code size (box_size * size)
            border: Border size in boxes
            error_correction: Error correction level ('L', 'M', 'Q', 'H')
            box_size: Size of each box in pixels
        
        Returns:
            Path to generated QR code image or None if failed
        """
        if not self.qrcode_available:
            logger.error("qrcode not available for QR code generation")
            return None
        
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                safe_data = "".join(c for c in data[:20] if c.isalnum() or c in (' ', '-', '_')).strip()
                safe_data = safe_data.replace(' ', '_')
                filename = f"qrcode_{safe_data}_{timestamp}.png"
            
            if not filename.endswith(('.png', '.jpg', '.jpeg')):
                filename += '.png'
            
            filepath = self.output_dir / filename
            
            # Create QR code
            qr = qrcode.QRCode(
                version=1,
                error_correction=getattr(qrcode.constants.ERROR_CORRECT, error_correction.upper(), qrcode.constants.ERROR_CORRECT_M),
                box_size=box_size,
                border=border,
            )
            qr.add_data(data)
            qr.make(fit=True)
            
            # Create image
            img = qr.make_image(fill_color="black", back_color="white")
            
            # Resize if needed
            if size != box_size:
                from PIL import Image
                img = img.resize((size * box_size, size * box_size), Image.Resampling.LANCZOS)
            
            # Save image
            img.save(str(filepath))
            
            logger.info(f"QR code generated: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generating QR code: {e}", exc_info=True)
            return None
    
    def generate_barcode(self,
                         data: str,
                         barcode_type: str = 'code128',
                         filename: str = None,
                         writer_options: Dict = None) -> Optional[str]:
        """
        Generate barcode image
        
        Args:
            data: Data to encode in barcode
            barcode_type: Barcode type ('ean13', 'ean8', 'code128', 'code39', 'upc', etc.)
            filename: Output filename (auto-generated if None)
            writer_options: Additional options for barcode writer
        
        Returns:
            Path to generated barcode image or None if failed
        """
        if not self.barcode_available:
            logger.error("python-barcode not available for barcode generation")
            return None
        
        try:
            # Get barcode class
            try:
                barcode_class = barcode.get_barcode_class(barcode_type)
            except Exception as e:
                logger.error(f"Invalid barcode type {barcode_type}: {e}")
                # Default to code128
                barcode_class = barcode.get_barcode_class('code128')
            
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                safe_data = "".join(c for c in data[:20] if c.isalnum() or c in (' ', '-', '_')).strip()
                safe_data = safe_data.replace(' ', '_')
                filename = f"barcode_{barcode_type}_{safe_data}_{timestamp}.png"
            
            if not filename.endswith(('.png', '.jpg', '.jpeg', '.svg')):
                filename += '.png'
            
            filepath = self.output_dir / filename
            
            # Create barcode
            code = barcode_class(data, writer=ImageWriter())
            
            # Default writer options
            default_options = {
                'module_width': 0.5,
                'module_height': 15.0,
                'quiet_zone': 6.5,
                'font_size': 10,
                'text_distance': 5.0,
                'background': 'white',
                'foreground': 'black',
            }
            
            # Merge with user options
            options = {**default_options, **(writer_options or {})}
            
            # Save barcode
            code.save(str(filepath).replace('.png', '').replace('.jpg', '').replace('.jpeg', '').replace('.svg', ''), options)
            
            # Get actual saved file (barcode library may add extension)
            saved_files = list(self.output_dir.glob(f"{Path(filepath).stem}*"))
            if saved_files:
                filepath = saved_files[0]
            
            logger.info(f"Barcode generated: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generating barcode: {e}", exc_info=True)
            return None
    
    def add_qr_code_to_pdf(self,
                          pdf_path: str,
                          qr_data: str,
                          position: Tuple[float, float] = None,
                          size: float = 1.0) -> Optional[str]:
        """
        Add QR code to existing PDF document
        
        Args:
            pdf_path: Path to PDF file
            qr_data: Data to encode in QR code
            position: (x, y) position in inches (None = bottom right)
            size: QR code size in inches
        
        Returns:
            Path to modified PDF or None if failed
        """
        if not self.pdf_available or not self.qrcode_available:
            logger.error("PDF or QR code generation not available")
            return None
        
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from PyPDF2 import PdfReader, PdfWriter
            import io
            
            # Generate QR code
            qr_path = self.generate_qr_code(qr_data, size=int(size * 72))  # Convert inches to points
            if not qr_path:
                return None
            
            # Read existing PDF
            reader = PdfReader(pdf_path)
            writer = PdfWriter()
            
            # Add QR code to each page
            for page_num, page in enumerate(reader.pages):
                # Create overlay with QR code
                packet = io.BytesIO()
                can = canvas.Canvas(packet, pagesize=letter)
                
                # Calculate position
                if position is None:
                    # Bottom right
                    x = letter[0] - (size * 72) - 20
                    y = 20
                else:
                    x = position[0] * 72
                    y = position[1] * 72
                
                # Draw QR code
                can.drawImage(qr_path, x, y, width=size * 72, height=size * 72)
                can.save()
                
                # Merge overlay
                packet.seek(0)
                overlay = PdfReader(packet)
                page.merge_page(overlay.pages[0])
                writer.add_page(page)
            
            # Save modified PDF
            output_path = Path(pdf_path).parent / f"{Path(pdf_path).stem}_with_qr.pdf"
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            # Clean up temp QR code
            if Path(qr_path).exists():
                Path(qr_path).unlink()
            
            logger.info(f"QR code added to PDF: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error adding QR code to PDF: {e}", exc_info=True)
            return None
    
    def add_barcode_to_pdf(self,
                          pdf_path: str,
                          barcode_data: str,
                          barcode_type: str = 'code128',
                          position: Tuple[float, float] = None,
                          size: Tuple[float, float] = (3.0, 1.0)) -> Optional[str]:
        """
        Add barcode to existing PDF document
        
        Args:
            pdf_path: Path to PDF file
            barcode_data: Data to encode in barcode
            barcode_type: Barcode type
            position: (x, y) position in inches (None = bottom right)
            size: (width, height) in inches
        
        Returns:
            Path to modified PDF or None if failed
        """
        if not self.pdf_available or not self.barcode_available:
            logger.error("PDF or barcode generation not available")
            return None
        
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from PyPDF2 import PdfReader, PdfWriter
            import io
            
            # Generate barcode
            barcode_path = self.generate_barcode(barcode_data, barcode_type)
            if not barcode_path:
                return None
            
            # Read existing PDF
            reader = PdfReader(pdf_path)
            writer = PdfWriter()
            
            # Add barcode to each page
            for page_num, page in enumerate(reader.pages):
                # Create overlay with barcode
                packet = io.BytesIO()
                can = canvas.Canvas(packet, pagesize=letter)
                
                # Calculate position
                if position is None:
                    # Bottom right
                    x = letter[0] - (size[0] * 72) - 20
                    y = 20
                else:
                    x = position[0] * 72
                    y = position[1] * 72
                
                # Draw barcode
                can.drawImage(barcode_path, x, y, width=size[0] * 72, height=size[1] * 72)
                can.save()
                
                # Merge overlay
                packet.seek(0)
                overlay = PdfReader(packet)
                page.merge_page(overlay.pages[0])
                writer.add_page(page)
            
            # Save modified PDF
            output_path = Path(pdf_path).parent / f"{Path(pdf_path).stem}_with_barcode.pdf"
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            # Clean up temp barcode
            if Path(barcode_path).exists():
                Path(barcode_path).unlink()
            
            logger.info(f"Barcode added to PDF: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error adding barcode to PDF: {e}", exc_info=True)
            return None


# Global instance
_document_generator_instance = None

def get_document_generator(output_dir: str = "generated_documents") -> DocumentGenerator:
    """Get or create global document generator instance"""
    global _document_generator_instance
    if _document_generator_instance is None:
        _document_generator_instance = DocumentGenerator(output_dir)
    return _document_generator_instance
